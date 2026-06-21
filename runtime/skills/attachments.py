from __future__ import annotations

import asyncio
from pathlib import Path
from typing import Any

from runtime.tool_registry import ToolRegistry, ToolSpec


async def extract_attachment_text(input_data: dict[str, Any], context: Any) -> dict[str, Any]:
    attachment_id = str(input_data.get("attachment_id") or "").strip()
    allowed_ids = set(getattr(context, "attachment_ids", ()) or ())
    if attachment_id not in allowed_ids:
        raise PermissionError("attachment is not available to the current conversation run")
    store = getattr(context, "attachment_store", None)
    if store is None:
        raise RuntimeError("attachment store is unavailable")
    record = store.get(attachment_id)
    if not record:
        raise KeyError(f"unknown attachment: {attachment_id}")
    max_chars = max(1, min(int(input_data.get("max_chars") or 100_000), 250_000))
    max_pages = max(1, min(int(input_data.get("max_pages") or 100), 500))
    text = await asyncio.to_thread(_extract_text, store.path_for(record), record.media_type, max_pages)
    truncated = len(text) > max_chars
    text = text[:max_chars]
    return {
        "type": "attachment_text",
        "attachment": record.to_public_dict(),
        "content": text,
        "text_chars": len(text),
        "truncated": truncated,
    }


def _extract_text(path: Path, media_type: str, max_pages: int) -> str:
    suffix = path.suffix.lower()
    if media_type == "application/pdf" or suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("pypdf is required to extract PDF attachments") from exc
        reader = PdfReader(str(path))
        return "\n\n".join((page.extract_text() or "") for page in reader.pages[:max_pages])
    if (
        media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or suffix == ".docx"
    ):
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("python-docx is required to extract Word attachments") from exc
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    content = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("attachment type does not support text extraction")


def register_attachment_tools(registry: ToolRegistry) -> None:
    registry.register(
        ToolSpec(
            id="attachment.extract_text",
            name="提取对话附件文本",
            description=(
                "从当前对话中用户上传的文本、PDF 或 Word 附件提取文字。附件是只读输入产物，不属于项目文件；"
                "使用消息附件目录中的 attachment_id 调用。"
            ),
            input_schema={
                "type": "object",
                "properties": {
                    "attachment_id": {"type": "string", "description": "消息附件目录中的附件 ID"},
                    "max_chars": {"type": "integer", "default": 100000, "description": "最多返回字符数"},
                    "max_pages": {"type": "integer", "default": 100, "description": "PDF 最多提取页数"},
                },
                "required": ["attachment_id"],
            },
            requires_confirmation=False,
            capability="attachment.user_input",
            artifacts=["attachment_text"],
            retry_safe=True,
            idempotent=True,
        ),
        extract_attachment_text,
    )
