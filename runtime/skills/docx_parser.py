"""
Word 文档智能解析器
移植自 aipython/core/plan_review.py，适配 local-intelligent-terminal 架构。

支持特性：
1. python-docx 正常提取（段落 + 表格）
2. 损坏/非标准 DOCX 文件自动降级（win32com → LibreOffice）
3. .doc 旧格式支持（win32com → LibreOffice 转换后提取）
4. 文本后处理清洗
"""
from __future__ import annotations

import asyncio
import logging
import os
import shutil
import subprocess
import tempfile
from concurrent.futures import ThreadPoolExecutor
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from .text_postprocessor import TextPostProcessor

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# 可选依赖检测
# ---------------------------------------------------------------------------

try:
    import win32com.client  # type: ignore
    import pythoncom  # type: ignore
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False


# ---------------------------------------------------------------------------
# 解析结果
# ---------------------------------------------------------------------------

@dataclass
class DocxParseResult:
    """Word 文档解析结果"""
    text: str = ""
    paragraph_count: int = 0
    table_count: int = 0
    headings: list[dict[str, Any]] = field(default_factory=list)
    strategy: str = "python-docx"  # python-docx | win32com | libreoffice
    warnings: list[str] = field(default_factory=list)


# ---------------------------------------------------------------------------
# 核心解析器
# ---------------------------------------------------------------------------

class DocxParser:
    """异步 Word 文档解析器，支持损坏文件降级和旧格式转换"""

    _executor = ThreadPoolExecutor(max_workers=2)
    _post_processor = TextPostProcessor()

    # 损坏文件特征关键词
    _CORRUPTED_KEYWORDS = ('null', 'archive', 'zipfile', 'bad magic', 'not a zip')

    # ------------------------------------------------------------------
    # 公开接口
    # ------------------------------------------------------------------

    async def parse(
        self,
        file_path: Path,
        extract_text: bool = True,
        extract_headings: bool = True,
    ) -> DocxParseResult:
        """主入口：智能解析 Word 文档。

        :param file_path: 文档路径（.docx 或 .doc）
        :param extract_text: 是否提取全文
        :param extract_headings: 是否提取标题大纲
        """
        result = DocxParseResult()
        suffix = file_path.suffix.lower()

        if suffix == '.docx':
            result = await self._parse_docx(file_path, extract_text, extract_headings)
        elif suffix == '.doc':
            result = await self._parse_doc(file_path, extract_text, extract_headings)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}，仅支持 .docx 和 .doc")

        # 后处理
        if result.text:
            result.text = self._post_processor.process(result.text)

        return result

    # ------------------------------------------------------------------
    # DOCX 解析（带容错降级）
    # ------------------------------------------------------------------

    async def _parse_docx(
        self,
        file_path: Path,
        extract_text: bool,
        extract_headings: bool,
    ) -> DocxParseResult:
        """解析 .docx 文件：先用 python-docx，失败时降级"""
        loop = asyncio.get_event_loop()

        try:
            return await loop.run_in_executor(
                self._executor,
                self._sync_extract_docx,
                file_path, extract_text, extract_headings,
            )
        except Exception as e:
            error_msg = str(e).lower()
            # 判断是否为损坏文件特征
            is_corrupted = any(kw in error_msg for kw in self._CORRUPTED_KEYWORDS)
            if is_corrupted:
                logger.warning(
                    f"python-docx 解析失败 ({str(e)[:80]})，降级转换: {file_path.name}"
                )
                return await loop.run_in_executor(
                    self._executor,
                    self._fallback_extract,
                    file_path, extract_text, extract_headings,
                )
            # 非损坏文件类错误，直接抛出
            raise

    # ------------------------------------------------------------------
    # DOC 旧格式
    # ------------------------------------------------------------------

    async def _parse_doc(
        self,
        file_path: Path,
        extract_text: bool,
        extract_headings: bool,
    ) -> DocxParseResult:
        """解析 .doc 文件：先尝试 python-docx（部分兼容），否则走降级"""
        loop = asyncio.get_event_loop()

        try:
            return await loop.run_in_executor(
                self._executor,
                self._sync_extract_docx,
                file_path, extract_text, extract_headings,
            )
        except Exception:
            logger.info(f".doc 格式 python-docx 不兼容，走降级转换: {file_path.name}")
            return await loop.run_in_executor(
                self._executor,
                self._fallback_extract,
                file_path, extract_text, extract_headings,
            )

    # ------------------------------------------------------------------
    # 同步提取（python-docx）
    # ------------------------------------------------------------------

    @staticmethod
    def _sync_extract_docx(
        file_path: Path,
        extract_text: bool,
        extract_headings: bool,
    ) -> DocxParseResult:
        """使用 python-docx 提取文档内容"""
        from docx import Document

        doc = Document(str(file_path))
        result = DocxParseResult(strategy="python-docx")

        paragraphs: list[str] = []
        para_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            para_count += 1

            if extract_text:
                paragraphs.append(text)

            if extract_headings:
                style_name = para.style.name if para.style else ""
                if style_name.lower().startswith("heading"):
                    result.headings.append({
                        "text": text,
                        "style": style_name,
                    })

        result.paragraph_count = para_count

        # 提取表格内容
        if extract_text and doc.tables:
            result.table_count = len(doc.tables)
            for table in doc.tables:
                for row in table.rows:
                    row_texts: list[str] = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_texts.append(cell_text)
                    if row_texts:
                        paragraphs.append(' | '.join(row_texts))

        if extract_text:
            result.text = '\n'.join(paragraphs)

        return result

    # ------------------------------------------------------------------
    # 降级方案：win32com → LibreOffice
    # ------------------------------------------------------------------

    def _fallback_extract(
        self,
        file_path: Path,
        extract_text: bool,
        extract_headings: bool,
    ) -> DocxParseResult:
        """降级提取：优先 win32com，备选 LibreOffice 转换"""
        # 方案 1: win32com（Windows + Office/WPS）
        if HAS_WIN32COM:
            try:
                docx_path = self._convert_via_win32com(file_path)
                if docx_path:
                    result = self._sync_extract_docx(
                        Path(docx_path), extract_text, extract_headings
                    )
                    result.strategy = "win32com"
                    # 清理临时文件
                    try:
                        os.unlink(docx_path)
                    except Exception:
                        pass
                    return result
            except Exception as e:
                logger.debug(f"win32com 转换失败: {e}")

        # 方案 2: LibreOffice 命令行
        soffice = self._find_libreoffice()
        if soffice:
            try:
                docx_path = self._convert_via_libreoffice(file_path, soffice)
                if docx_path:
                    result = self._sync_extract_docx(
                        Path(docx_path), extract_text, extract_headings
                    )
                    result.strategy = "libreoffice"
                    # 清理临时文件
                    try:
                        os.unlink(docx_path)
                    except Exception:
                        pass
                    return result
            except Exception as e:
                logger.debug(f"LibreOffice 转换失败: {e}")

        # 全部失败
        result = DocxParseResult()
        result.warnings.append(
            "文件解析失败。python-docx 无法处理此文件（可能是损坏/非标准格式），"
            "且 win32com 和 LibreOffice 均不可用。"
            "建议：1. 安装 pywin32 + Office/WPS；2. 安装 LibreOffice；"
            "3. 手动用 Word 另存为标准 .docx 格式。"
        )
        return result

    # ------------------------------------------------------------------
    # 转换工具
    # ------------------------------------------------------------------

    @staticmethod
    def _convert_via_win32com(file_path: Path) -> str | None:
        """使用 Windows COM 接口转换为 docx"""
        import pythoncom
        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            doc = word.Documents.Open(str(file_path.absolute()))

            tmpdir = tempfile.mkdtemp(prefix="docx_convert_")
            docx_path = os.path.join(tmpdir, f"{file_path.stem}.docx")
            # FileFormat=16 = wdFormatXMLDocument (.docx)
            doc.SaveAs2(docx_path, FileFormat=16)
            doc.Close(False)
            word.Quit()
            pythoncom.CoUninitialize()
            return docx_path
        except Exception as e:
            if doc:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()
            raise

    @staticmethod
    def _find_libreoffice() -> str | None:
        """查找系统中的 LibreOffice 可执行文件"""
        candidates = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "libreoffice",
            "soffice",
        ]
        for path in candidates:
            if shutil.which(path) or Path(path).exists():
                return path
        return None

    @staticmethod
    def _convert_via_libreoffice(file_path: Path, soffice: str) -> str | None:
        """使用 LibreOffice 命令行转换为 docx"""
        with tempfile.TemporaryDirectory(prefix="lo_convert_") as tmpdir:
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "docx", "--outdir", tmpdir, str(file_path)],
                capture_output=True,
                timeout=60,
            )
            if result.returncode != 0:
                stderr = result.stderr.decode('utf-8', errors='ignore')
                raise RuntimeError(f"LibreOffice 转换失败: {stderr[:200]}")

            docx_files = list(Path(tmpdir).glob("*.docx"))
            if not docx_files:
                raise RuntimeError("LibreOffice 转换后未生成 .docx 文件")

            # 将文件移出临时目录（因为 TemporaryDirectory 会删除）
            final_path = tempfile.mktemp(suffix=".docx", prefix="lo_out_")
            shutil.copy2(str(docx_files[0]), final_path)
            return final_path
