from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
import sys
import types

import pytest

from runtime.security import PathGuard
from runtime.skills import pdf_parser
from runtime.skills import document as document_tools
from runtime.skills.document import export_docx, extract_pdf_to_docx, translate_docx


@dataclass
class FakeContext:
    path_guard: PathGuard
    settings: object | None = None
    logs: list[tuple[str, str, dict | None]] = field(default_factory=list)

    def log(self, level: str, message: str, data: dict | None = None) -> None:
        self.logs.append((level, message, data))


class FakeSettings:
    def get_default_model(self) -> str:
        return "fake-model"


@pytest.mark.asyncio
async def test_extract_pdf_to_docx_writes_docx_with_parsed_text(tmp_path: Path, monkeypatch) -> None:
    class FakePDFParser:
        async def parse(self, file_path: Path, max_pages: int = 0, context=None):
            return pdf_parser.ParseResult(
                text="第一段 PDF 文本。\n\nSecond paragraph.",
                total_pages=2,
                pages_parsed=2,
                strategy="fake",
            )

    monkeypatch.setattr(pdf_parser, "PDFParser", FakePDFParser)
    pdf_path = tmp_path / "sample.pdf"
    pdf_path.write_bytes(b"%PDF-1.4\n")
    output_path = tmp_path / "sample.docx"

    result = await extract_pdf_to_docx(
        {
            "path": str(pdf_path),
            "output_path": str(output_path),
            "title": "PDF 提取测试",
        },
        FakeContext(PathGuard([tmp_path])),
    )

    assert result["path"] == str(output_path.resolve())
    assert result["source_path"] == str(pdf_path.resolve())
    assert result["pages_parsed"] == 2
    assert output_path.exists()
    assert output_path.stat().st_size > 0


@pytest.mark.asyncio
async def test_extract_pdf_to_docx_text_with_images_preserves_ordered_image_blocks(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    image_module = pytest.importorskip("PIL.Image")
    from docx import Document

    pdf_path = tmp_path / "with-image.pdf"
    output_path = tmp_path / "with-image.docx"

    image = image_module.new("RGB", (40, 30), color=(220, 40, 40))
    image_bytes = BytesIO()
    image.save(image_bytes, format="PNG")

    pdf = fitz.open()
    page = pdf.new_page(width=300, height=300)
    page.insert_text((40, 50), "Before image")
    page.insert_image(fitz.Rect(40, 80, 140, 150), stream=image_bytes.getvalue())
    page.insert_text((40, 190), "After image")
    pdf.save(str(pdf_path))
    pdf.close()

    context = FakeContext(PathGuard([tmp_path]))
    result = await extract_pdf_to_docx(
        {
            "path": str(pdf_path),
            "output_path": str(output_path),
            "mode": "text_with_images",
            "title": "PDF with image",
        },
        context,
    )

    doc = Document(str(output_path))
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

    assert result["mode"] == "text_with_images"
    assert result["strategy"] == "pymupdf_blocks"
    assert result["image_count"] >= 1
    assert result["text_block_count"] >= 1
    assert result["file_size"] > 0
    assert "Before image" in text
    assert "After image" in text
    assert len(doc.inline_shapes) >= 1
    assert any(message.startswith("pdf page converted ") for _level, message, _data in context.logs)
    assert any(data and data.get("kind") == "pdf_to_docx" for _level, _message, data in context.logs)


@pytest.mark.asyncio
async def test_export_docx_returns_content_size_metadata(tmp_path: Path) -> None:
    output_path = tmp_path / "report.docx"

    result = await export_docx(
        {
            "path": str(output_path),
            "title": "Report",
            "content": "# Report\n\nFirst paragraph.\n\nSecond paragraph.",
        },
        FakeContext(PathGuard([tmp_path])),
    )

    assert result["path"] == str(output_path.resolve())
    assert result["content_chars"] > 0
    assert result["paragraph_count"] >= 2
    assert result["nonempty_paragraph_count"] >= 2
    assert result["file_size"] > 0


@pytest.mark.asyncio
async def test_translate_docx_completes_all_paragraphs(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        return [f"ZH:{text}" for text in kwargs["texts"]]

    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(source_path))
    output_path = tmp_path / "source_zh.docx"

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(output_path),
            "target_language": "zh-CN",
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    assert result["complete"] is True
    assert result["engine"] == "model"
    assert result["model"] == "fake-model"
    assert result["translation_profile"] == "balanced"
    assert result["max_chars_per_chunk"] == 4000
    assert result["max_chars_per_batch"] == 12000
    assert result["max_paragraphs_per_batch"] == 8
    assert result["batch_timeout"] == 240
    assert result["translated_paragraph_count"] == 2
    assert result["failed_paragraph_count"] == 0
    assert "error" not in result
    assert output_path.exists()


@pytest.mark.asyncio
async def test_translate_docx_fast_profile_uses_larger_batches(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        return [f"ZH:{text}" for text in kwargs["texts"]]

    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.save(str(source_path))

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(tmp_path / "source_zh.docx"),
            "target_language": "zh-CN",
            "translation_profile": "fast",
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    assert result["complete"] is True
    assert result["translation_profile"] == "fast"
    assert result["max_chars_per_chunk"] == 6000
    assert result["max_chars_per_batch"] == 24000
    assert result["max_paragraphs_per_batch"] == 16
    assert result["batch_timeout"] == 300


@pytest.mark.asyncio
async def test_translate_docx_logs_progress_during_batch_fallback(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        texts = kwargs["texts"]
        if len(texts) > 1:
            raise RuntimeError("batch failed")
        return [f"ZH:{texts[0]}"]

    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    doc.save(str(source_path))
    context = FakeContext(PathGuard([tmp_path]), settings=FakeSettings())

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(tmp_path / "fallback.docx"),
            "max_paragraphs_per_batch": 3,
            "max_chars_per_batch": 10000,
        },
        context,
    )

    fallback_logs = [
        data
        for _level, message, data in context.logs
        if message.startswith("translation progress ") and data and data.get("fallback")
    ]

    assert result["complete"] is True
    assert result["translated_paragraph_count"] == 3
    assert len(fallback_logs) == 3
    assert [log["paragraph_index"] for log in fallback_logs] == [1, 2, 3]
    assert fallback_logs[-1]["source_chars_done"] == result["source_chars_done"]


@pytest.mark.asyncio
async def test_translate_docx_marks_limited_output_partial(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        return [f"ZH:{text}" for text in kwargs["texts"]]

    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(source_path))

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(tmp_path / "partial.docx"),
            "max_paragraphs": 1,
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    assert result["complete"] is False
    assert result["status"] == "partial"
    assert result["partial_resumable"] is False
    assert result["translated_paragraph_count"] == 1
    assert result["stopped_reason"] == "max_paragraphs_reached:1"


@pytest.mark.asyncio
async def test_translate_docx_can_resume_from_checkpoint(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        return [f"ZH:{text}" for text in kwargs["texts"]]

    ticks = iter([0, 0, 1900])
    monkeypatch.setattr(document_tools.time, "monotonic", lambda: next(ticks, 1900))
    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    output_path = tmp_path / "source_zh.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.add_paragraph("Third paragraph.")
    doc.save(str(source_path))

    first = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(output_path),
            "max_seconds": 1800,
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    assert first["status"] == "partial_resumable"
    assert first["partial_resumable"] is True
    assert first["translated_paragraph_count"] == 1
    assert Path(first["manifest_path"]).exists()

    monkeypatch.setattr(document_tools.time, "monotonic", lambda: 0)
    second = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(output_path),
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    resumed = Document(str(output_path))
    text = "\n".join(paragraph.text for paragraph in resumed.paragraphs)

    assert second["complete"] is True
    assert second["status"] == "success"
    assert second["resumed_from_checkpoint"] is True
    assert second["translated_paragraph_count"] == 3
    assert "ZH:First paragraph." in text
    assert "ZH:Second paragraph." in text
    assert "ZH:Third paragraph." in text


@pytest.mark.asyncio
async def test_translate_docx_raises_low_full_document_timeout(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    async def fake_translate_batch(**kwargs):
        return [f"ZH:{text}" for text in kwargs["texts"]]

    monkeypatch.setattr(document_tools, "_translate_text_batch_with_model", fake_translate_batch)
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(source_path))

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(tmp_path / "raised_timeout.docx"),
            "max_seconds": 110,
        },
        FakeContext(PathGuard([tmp_path]), settings=FakeSettings()),
    )

    assert result["complete"] is True
    assert result["max_seconds"] == 1800
    assert result["translated_paragraph_count"] == 2


@pytest.mark.asyncio
async def test_translate_docx_can_use_google_engine_explicitly(tmp_path: Path, monkeypatch) -> None:
    from docx import Document

    class FakeTranslator:
        def __init__(self, source: str, target: str) -> None:
            self.source = source
            self.target = target

        def translate(self, text: str) -> str:
            return f"ZH:{text}"

    monkeypatch.setitem(
        sys.modules,
        "deep_translator",
        types.SimpleNamespace(GoogleTranslator=FakeTranslator),
    )
    source_path = tmp_path / "source.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.save(str(source_path))
    output_path = tmp_path / "source_google.docx"

    result = await translate_docx(
        {
            "path": str(source_path),
            "output_path": str(output_path),
            "engine": "google",
        },
        FakeContext(PathGuard([tmp_path])),
    )

    assert result["complete"] is True
    assert result["engine"] == "google"
    assert result["translated_paragraph_count"] == 1
    assert output_path.exists()
